from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_readme():
    doc = Document()
    
    # Title
    title = doc.add_heading('DataSense AI', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Multi-Agent AI Data Analysis Pipeline & Dashboard')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Overview
    doc.add_heading('Overview', level=1)
    doc.add_paragraph('DataSense AI is a premium SaaS application that transforms natural language queries and raw datasets (CSV, Excel, JSON, Parquet, or SQL) into interactive dashboards and actionable business insights. It leverages a powerful 6-agent AI pipeline driven by Google Gemini and a robust Next.js frontend with pixel-perfect alignment and a modern design system.')
    
    # Features
    doc.add_heading('Key Features', level=1)
    
    features = [
        "Natural Language to Dashboard: Type a question, and the multi-agent pipeline generates the corresponding data visualizations.",
        "Universal Data Connectors: Drag-and-drop CSV, Excel, JSON, Parquet, or connect directly to a SQL database.",
        "AI-Generated Insights: Receive automatic, color-coded business insights summarizing findings from your data.",
        "Self-Correcting 6-Agent Pipeline: Includes ambiguity detection, planning, execution, and self-correction to ensure accurate data extraction and charting.",
        "Premium UI/UX: Fully redesigned with Framer Motion animations, a strict 8px grid, deep dark mode (#080C14), and Plus Jakarta Sans typography."
    ]
    for feat in features:
        doc.add_paragraph(feat, style='List Bullet')
        
    # Tech Stack
    doc.add_heading('Technology Stack', level=1)
    techs = [
        "Frontend: Next.js 14, React, Tailwind CSS (custom design system), Framer Motion, Lucide React, Vega-Lite (Charts).",
        "Backend: FastAPI, Python, DuckDB, Pandas.",
        "AI / LLM: Google Gemini API (gemini-2.0-flash / gemini-2.0-flash-lite) orchestrated via custom LangChain-style patterns."
    ]
    for tech in techs:
        doc.add_paragraph(tech, style='List Bullet')
        
    # Setup Instructions
    doc.add_heading('Getting Started', level=1)
    
    doc.add_heading('1. Backend Setup', level=2)
    doc.add_paragraph('1. Navigate to the backend directory: cd backend')
    doc.add_paragraph('2. Create and activate a virtual environment: python -m venv venv')
    doc.add_paragraph('3. Install dependencies: pip install -r requirements.txt')
    doc.add_paragraph('4. Configure your .env file with a GEMINI_API_KEY')
    doc.add_paragraph('5. Run the backend: python -m uvicorn main:app --reload --port 8000')
    
    doc.add_heading('2. Frontend Setup', level=2)
    doc.add_paragraph('1. Navigate to the frontend directory: cd frontend')
    doc.add_paragraph('2. Install dependencies: npm install')
    doc.add_paragraph('3. Run the development server: npm run dev')
    doc.add_paragraph('4. Open http://localhost:3000 in your browser')
    
    # File Structure
    doc.add_heading('Project Structure', level=1)
    struct = [
        "frontend/app/page.js: Main application logic and state management.",
        "frontend/app/components/: Modular UI components (Navbar, Sidebar, ChatInterface, DashboardView, InsightsPanel, etc.).",
        "frontend/app/globals.css: Global design tokens and utilities.",
        "backend/main.py: FastAPI entry point.",
        "backend/app/agents/: The multi-agent orchestrator and task flow logic."
    ]
    for s in struct:
        doc.add_paragraph(s, style='List Bullet')

    doc.save(r'd:\Dashboard\DV PROJECT\README.docx')

if __name__ == '__main__':
    create_readme()
